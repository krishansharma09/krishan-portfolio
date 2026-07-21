import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

def create_resume_docx(output_path):
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    styles = doc.styles
    style_normal = styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run("Krishan Sharma")
    run_name.font.size = Pt(24)
    run_name.font.bold = True
    run_name.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    p_role = doc.add_paragraph()
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_role.paragraph_format.space_after = Pt(6)
    run_role = p_role.add_run("Machine Learning Engineer")
    run_role.font.size = Pt(13)
    run_role.font.bold = True
    run_role.font.color.rgb = RGBColor(0x1E, 0x56, 0xA0)

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(14)
    run_c = p_contact.add_run(
        "krishansharma995060@gmail.com  •  +91 86191 64592  •  LinkedIn: /in/krishan-sharma-50a89635b\n"
        "GitHub: github.com/krishansharma09  •  Jaipur, Rajasthan"
    )
    run_c.font.size = Pt(9.5)
    run_c.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title.upper())
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x56, 0xA0)
        return p

    # Professional Summary
    add_section_header("Professional Summary")
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_after = Pt(10)
    p_sum.paragraph_format.line_spacing = 1.15
    p_sum.add_run(
        "Results-driven Machine Learning Engineer with 7 months of hands-on internship experience at Dotsquares Technologies, "
        "specializing in building AI-powered applications using Python, LangChain, RAG pipelines, and Generative AI frameworks. Proven "
        "ability to design and deploy NLP systems, RESTful APIs, and document intelligence solutions. Passionate about leveraging cutting-edge "
        "LLMs and vector databases to solve real-world problems at scale."
    )

    # Work Experience
    add_section_header("Work Experience")
    p_exp = doc.add_paragraph()
    p_exp.paragraph_format.space_after = Pt(4)
    r1 = p_exp.add_run("Machine Learning Engineer Intern")
    r1.bold = True
    p_exp.add_run(" | ")
    r2 = p_exp.add_run("Dotsquares Technologies Pvt. Ltd.")
    r2.bold = True
    r2.font.color.rgb = RGBColor(0x1E, 0x56, 0xA0)
    p_exp.add_run(" | India          ")
    r3 = p_exp.add_run("Aug 2024 – Feb 2025")
    r3.italic = True

    bullets_exp = [
        "Designed and deployed NLP pipelines for text classification, entity recognition, and semantic similarity, achieving over 90% model accuracy on internal benchmarks.",
        "Built and integrated RESTful APIs using Flask and FastAPI to serve ML model inference endpoints, reducing average response latency by 35% through async processing.",
        "Developed and fine-tuned machine learning models using Scikit-learn for business use cases including predictive analytics and document categorization.",
        "Implemented LLM-powered features using LangChain and Gemini AI, enabling intelligent document Q&A and automated text summarization workflows.",
        "Collaborated with cross-functional teams in an Agile environment to deliver 3 production-ready AI features within sprint timelines.",
        "Optimized PostgreSQL database schemas and integrated FAISS vector stores for efficient similarity search across large-scale text corpora."
    ]
    for b in bullets_exp:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.line_spacing = 1.1
        bp.add_run(b)

    # Projects
    add_section_header("Projects")

    # Project 1
    p_pj1 = doc.add_paragraph()
    p_pj1.paragraph_format.space_before = Pt(6)
    p_pj1.paragraph_format.space_after = Pt(3)
    r_title = p_pj1.add_run("AI Chatbot with Gemini + PostgreSQL + FAISS")
    r_title.bold = True
    r_tech = p_pj1.add_run(" — Python · LangChain · Gemini AI · FAISS · PostgreSQL · FastAPI")
    r_tech.font.color.rgb = RGBColor(0x1E, 0x56, 0xA0)
    r_tech.font.size = Pt(9.5)

    bullets_pj1 = [
        "Architected a production-grade conversational AI chatbot leveraging Google Gemini LLM with LangChain orchestration for multi-turn dialogue management.",
        "Integrated FAISS vector store for sub-second semantic search across 10,000+ document embeddings, enabling context-aware response generation.",
        "Built a FastAPI backend with PostgreSQL for persistent chat history storage, session management, and user authentication — supporting 50+ concurrent users.",
        "Achieved 92% user satisfaction score in internal testing by combining dense retrieval with Gemini's generative reasoning capabilities."
    ]
    for b in bullets_pj1:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.line_spacing = 1.1
        bp.add_run(b)

    # Project 2
    p_pj2 = doc.add_paragraph()
    p_pj2.paragraph_format.space_before = Pt(6)
    p_pj2.paragraph_format.space_after = Pt(3)
    r_title = p_pj2.add_run("RAG-Based Document Chatbot")
    r_title.bold = True
    r_tech = p_pj2.add_run(" — Python · LangChain · FAISS · Streamlit · OpenAI / Gemini")
    r_tech.font.color.rgb = RGBColor(0x1E, 0x56, 0xA0)
    r_tech.font.size = Pt(9.5)

    bullets_pj2 = [
        "Built an end-to-end Retrieval-Augmented Generation (RAG) pipeline allowing users to upload PDFs and query documents in natural language with high precision.",
        "Implemented document chunking, embedding generation, and vector indexing with FAISS, reducing irrelevant context injection by 40% vs. naive prompting.",
        "Deployed an interactive Streamlit UI enabling non-technical users to interact with AI-powered document search without any coding knowledge.",
        "Optimized prompt templates and retrieval chain logic to maintain factual accuracy and minimize hallucination in generated responses."
    ]
    for b in bullets_pj2:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.line_spacing = 1.1
        bp.add_run(b)

    # Project 3
    p_pj3 = doc.add_paragraph()
    p_pj3.paragraph_format.space_before = Pt(6)
    p_pj3.paragraph_format.space_after = Pt(3)
    r_title = p_pj3.add_run("Binance Futures Trading Bot")
    r_title.bold = True
    r_tech = p_pj3.add_run(" — Python · Binance API · WebSocket · Pandas · NumPy")
    r_tech.font.color.rgb = RGBColor(0x1E, 0x56, 0xA0)
    r_tech.font.size = Pt(9.5)

    bullets_pj3 = [
        "Developed an automated cryptocurrency futures trading bot using Python and Binance API, capable of executing real-time buy/sell orders based on custom algorithmic strategies.",
        "Implemented WebSocket-based live market data streaming for sub-second price feed ingestion, enabling low-latency trading signal generation across multiple trading pairs.",
        "Designed a modular strategy engine supporting technical indicators (RSI, MACD, Bollinger Bands) with configurable risk management including stop-loss and take-profit automation.",
        "Built a backtesting framework using Pandas and NumPy to simulate strategies on historical OHLCV data, optimizing entry/exit parameters before live deployment."
    ]
    for b in bullets_pj3:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.line_spacing = 1.1
        bp.add_run(b)

    # Technical Skills
    add_section_header("Technical Skills")
    skills = [
        ("Languages", "Python (Advanced), SQL, Bash"),
        ("AI / ML", "Machine Learning, NLP, Deep Learning, Generative AI, Prompt Engineering"),
        ("Frameworks", "LangChain, Scikit-learn, TensorFlow (basics), Hugging Face Transformers"),
        ("Backend", "FastAPI, Flask, RESTful APIs, JWT Authentication, Pydantic"),
        ("Databases", "PostgreSQL, FAISS (Vector Store), SQLAlchemy ORM"),
        ("Tools & Platforms", "Streamlit, Gemini AI, OpenAI API, Git, GitHub, Docker (basics), Postman"),
        ("Methodologies", "Agile / Scrum, RAG Architecture, API Design, Model Deployment")
    ]
    for category, items in skills:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(2)
        r_cat = sp.add_run(f"{category}: ")
        r_cat.bold = True
        sp.add_run(items)

    # Education
    add_section_header("Education")
    
    p_ed1 = doc.add_paragraph()
    p_ed1.paragraph_format.space_after = Pt(2)
    r_deg1 = p_ed1.add_run("Master of Computer Applications (MCA)")
    r_deg1.bold = True
    p_ed1.add_run("          Aug 2025 – Present\n")
    p_ed1.add_run("JECRC University, Sitapura, Jaipur  |  ")
    r_st = p_ed1.add_run("Currently Pursuing")
    r_st.font.color.rgb = RGBColor(0x1E, 0x56, 0xA0)

    p_ed2 = doc.add_paragraph()
    p_ed2.paragraph_format.space_before = Pt(4)
    p_ed2.paragraph_format.space_after = Pt(2)
    r_deg2 = p_ed2.add_run("Bachelor of Computer Applications (BCA)")
    r_deg2.bold = True
    p_ed2.add_run("          Aug 2022 – Aug 2025\n")
    p_ed2.add_run("University of Rajasthan, Jaipur")

    doc.save(output_path)
    print(f"DOCX created at: {output_path}")

def create_resume_pdf(output_path):
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()

    primary_color = colors.HexColor('#1E56A0')
    dark_color = colors.HexColor('#111111')
    gray_color = colors.HexColor('#444444')

    name_style = ParagraphStyle(
        'NameStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=22,
        leading=26,
        alignment=TA_CENTER,
        textColor=dark_color,
        spaceAfter=2
    )

    role_style = ParagraphStyle(
        'RoleStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=15,
        alignment=TA_CENTER,
        textColor=primary_color,
        spaceAfter=6
    )

    contact_style = ParagraphStyle(
        'ContactStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        alignment=TA_CENTER,
        textColor=gray_color,
        spaceAfter=10
    )

    section_style = ParagraphStyle(
        'SectionStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=primary_color,
        spaceBefore=10,
        spaceAfter=4
    )

    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=dark_color,
        spaceAfter=6
    )

    bullet_style = ParagraphStyle(
        'BulletStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12.5,
        textColor=dark_color
    )

    story = []

    # Header
    story.append(Paragraph("Krishan Sharma", name_style))
    story.append(Paragraph("Machine Learning Engineer", role_style))
    story.append(Paragraph(
        "krishansharma995060@gmail.com &nbsp;&bull;&nbsp; +91 86191 64592 &nbsp;&bull;&nbsp; LinkedIn: /in/krishan-sharma-50a89635b &nbsp;&bull;&nbsp; GitHub: github.com/krishansharma09 &nbsp;&bull;&nbsp; Jaipur, Rajasthan",
        contact_style
    ))
    story.append(HRFlowable(width="100%", thickness=1, color=primary_color, spaceAfter=8, spaceBefore=0))

    # Summary
    story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#CCCCCC'), spaceAfter=6, spaceBefore=0))
    story.append(Paragraph(
        "Results-driven Machine Learning Engineer with 7 months of hands-on internship experience at Dotsquares Technologies, "
        "specializing in building AI-powered applications using Python, LangChain, RAG pipelines, and Generative AI frameworks. Proven "
        "ability to design and deploy NLP systems, RESTful APIs, and document intelligence solutions. Passionate about leveraging cutting-edge "
        "LLMs and vector databases to solve real-world problems at scale.",
        body_style
    ))

    # Experience
    story.append(Paragraph("WORK EXPERIENCE", section_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#CCCCCC'), spaceAfter=6, spaceBefore=0))
    story.append(Paragraph(
        "<b>Machine Learning Engineer Intern</b> &nbsp;|&nbsp; <font color='#1E56A0'><b>Dotsquares Technologies Pvt. Ltd.</b></font> &nbsp;|&nbsp; India &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; <i>Aug 2024 – Feb 2025</i>",
        body_style
    ))

    bullets_exp = [
        "Designed and deployed NLP pipelines for text classification, entity recognition, and semantic similarity, achieving over 90% model accuracy on internal benchmarks.",
        "Built and integrated RESTful APIs using Flask and FastAPI to serve ML model inference endpoints, reducing average response latency by 35% through async processing.",
        "Developed and fine-tuned machine learning models using Scikit-learn for business use cases including predictive analytics and document categorization.",
        "Implemented LLM-powered features using LangChain and Gemini AI, enabling intelligent document Q&A and automated text summarization workflows.",
        "Collaborated with cross-functional teams in an Agile environment to deliver 3 production-ready AI features within sprint timelines.",
        "Optimized PostgreSQL database schemas and integrated FAISS vector stores for efficient similarity search across large-scale text corpora."
    ]
    exp_items = [ListItem(Paragraph(b, bullet_style), leftIndent=12, bulletColor=dark_color) for b in bullets_exp]
    story.append(ListFlowable(exp_items, bulletType='bullet', start='square', spaceAfter=8))

    # Projects
    story.append(Paragraph("PROJECTS", section_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#CCCCCC'), spaceAfter=6, spaceBefore=0))

    # Project 1
    story.append(Paragraph("<b>AI Chatbot with Gemini + PostgreSQL + FAISS</b> &nbsp;—&nbsp; <font color='#1E56A0'>Python &middot; LangChain &middot; Gemini AI &middot; FAISS &middot; PostgreSQL &middot; FastAPI</font>", body_style))
    p1_bullets = [
        "Architected a production-grade conversational AI chatbot leveraging Google Gemini LLM with LangChain orchestration for multi-turn dialogue management.",
        "Integrated FAISS vector store for sub-second semantic search across 10,000+ document embeddings, enabling context-aware response generation.",
        "Built a FastAPI backend with PostgreSQL for persistent chat history storage, session management, and user authentication — supporting 50+ concurrent users.",
        "Achieved 92% user satisfaction score in internal testing by combining dense retrieval with Gemini's generative reasoning capabilities."
    ]
    story.append(ListFlowable([ListItem(Paragraph(b, bullet_style), leftIndent=12) for b in p1_bullets], bulletType='bullet', spaceAfter=6))

    # Project 2
    story.append(Paragraph("<b>RAG-Based Document Chatbot</b> &nbsp;—&nbsp; <font color='#1E56A0'>Python &middot; LangChain &middot; FAISS &middot; Streamlit &middot; OpenAI / Gemini</font>", body_style))
    p2_bullets = [
        "Built an end-to-end Retrieval-Augmented Generation (RAG) pipeline allowing users to upload PDFs and query documents in natural language with high precision.",
        "Implemented document chunking, embedding generation, and vector indexing with FAISS, reducing irrelevant context injection by 40% vs. naive prompting.",
        "Deployed an interactive Streamlit UI enabling non-technical users to interact with AI-powered document search without any coding knowledge.",
        "Optimized prompt templates and retrieval chain logic to maintain factual accuracy and minimize hallucination in generated responses."
    ]
    story.append(ListFlowable([ListItem(Paragraph(b, bullet_style), leftIndent=12) for b in p2_bullets], bulletType='bullet', spaceAfter=6))

    # Project 3
    story.append(Paragraph("<b>Binance Futures Trading Bot</b> &nbsp;—&nbsp; <font color='#1E56A0'>Python &middot; Binance API &middot; WebSocket &middot; Pandas &middot; NumPy</font>", body_style))
    p3_bullets = [
        "Developed an automated cryptocurrency futures trading bot using Python and Binance API, capable of executing real-time buy/sell orders based on custom algorithmic strategies.",
        "Implemented WebSocket-based live market data streaming for sub-second price feed ingestion, enabling low-latency trading signal generation across multiple trading pairs.",
        "Designed a modular strategy engine supporting technical indicators (RSI, MACD, Bollinger Bands) with configurable risk management including stop-loss and take-profit automation.",
        "Built a backtesting framework using Pandas and NumPy to simulate strategies on historical OHLCV data, optimizing entry/exit parameters before live deployment."
    ]
    story.append(ListFlowable([ListItem(Paragraph(b, bullet_style), leftIndent=12) for b in p3_bullets], bulletType='bullet', spaceAfter=8))

    # Technical Skills
    story.append(Paragraph("TECHNICAL SKILLS", section_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#CCCCCC'), spaceAfter=6, spaceBefore=0))
    skills = [
        ("Languages", "Python (Advanced), SQL, Bash"),
        ("AI / ML", "Machine Learning, NLP, Deep Learning, Generative AI, Prompt Engineering"),
        ("Frameworks", "LangChain, Scikit-learn, TensorFlow (basics), Hugging Face Transformers"),
        ("Backend", "FastAPI, Flask, RESTful APIs, JWT Authentication, Pydantic"),
        ("Databases", "PostgreSQL, FAISS (Vector Store), SQLAlchemy ORM"),
        ("Tools & Platforms", "Streamlit, Gemini AI, OpenAI API, Git, GitHub, Docker (basics), Postman"),
        ("Methodologies", "Agile / Scrum, RAG Architecture, API Design, Model Deployment")
    ]
    for category, items in skills:
        story.append(Paragraph(f"<b>{category}:</b> {items}", ParagraphStyle('Sk', parent=body_style, spaceAfter=2)))

    story.append(Spacer(1, 4))

    # Education
    story.append(Paragraph("EDUCATION", section_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#CCCCCC'), spaceAfter=6, spaceBefore=0))
    story.append(Paragraph("<b>Master of Computer Applications (MCA)</b> &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; <i>Aug 2025 – Present</i>", body_style))
    story.append(Paragraph("JECRC University, Sitapura, Jaipur &nbsp;|&nbsp; <font color='#1E56A0'>Currently Pursuing</font>", ParagraphStyle('Ed1', parent=body_style, spaceAfter=4)))
    story.append(Paragraph("<b>Bachelor of Computer Applications (BCA)</b> &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; <i>Aug 2022 – Aug 2025</i>", body_style))
    story.append(Paragraph("University of Rajasthan, Jaipur", ParagraphStyle('Ed2', parent=body_style, spaceAfter=4)))

    doc.build(story)
    print(f"PDF created at: {output_path}")

if __name__ == '__main__':
    target_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), 'app', 'static', 'assets'))
    os.makedirs(target_dir, exist_ok=True)
    docx_path = os.path.join(target_dir, 'Krishan_Sharma_Resume.docx')
    pdf_path = os.path.join(target_dir, 'Krishan_Sharma_Resume.pdf')
    create_resume_docx(docx_path)
    create_resume_pdf(pdf_path)
